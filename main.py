import sys
import os
import PyPDF2
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox
import os



# Função que lê todas as páginas do PDF
def ler_ocorrencia(pdf_path):
    texto_completo = ""
    try:
        with open(pdf_path, "rb") as pdf_file:
            leitor_pdf = PyPDF2.PdfReader(pdf_file)
            num_paginas = len(leitor_pdf.pages)
            for i in range(num_paginas):
                pagina = leitor_pdf.pages[i]
                texto_completo += pagina.extract_text() + "\n"
    except Exception as e:
        print(f"Erro ao ler PDF: {str(e)}")
    return texto_completo



def extrair_dados_necessarios(dados_pdf):
    rai = "RAI não encontrado"
    odin = "Odin não encontrado"
    requisitante = "Requisitante não encontrado"
    delegacia_requisitante = "Delegacia Requisitante não encontrada"
    unidade_afeta = "Unidade Afeta não encontrada"
    pessoa_envolvida = "Pessoa envolvida não encontrada"
    relato_pc = ""
    endereco = "Endereço não encontrado" 
    linhas = dados_pdf.split('\n')
    coletando_relato = False
    coletando_pessoa_envolvida = False

    for linha in linhas:
        # Converter a linha para minúsculas para facilitar a busca
        linha_lower = linha.lower()
        if "endereço:" in linha_lower:
            endereco = linha.split("Endereço:")[1].strip()

        if "rai:" in linha_lower:
            partes = linha.split("RAI:")
            if len(partes) > 1:
                rai = partes[1].strip()

        if "ocorrência nº" in linha_lower:
            partes = linha.split("Ocorrência nº")
            if len(partes) > 1:
                odin = partes[1].strip()

        if "requisitante:" in linha_lower:
            partes = linha.split("Requisitante:")
            if len(partes) > 1:
                requisitante = partes[1].strip()

        # Captura a primeira "Unidade" como Delegacia Requisitante
        if "unidade" in linha_lower and "afeta" not in linha_lower:
            partes = linha.split("Unidade:")
            if len(partes) > 1:
                delegacia_requisitante = partes[1].strip()

        # Captura a "Unidade Afeta"
        if "afeta" in linha_lower:
            partes = linha.split("Unidade Afeta:")
            if len(partes) > 1:
                unidade_afeta = partes[1].strip()

        # Iniciar a captura de pessoas envolvidas quando encontrar "Pessoas Envolvidas"
        if "pessoas envolvidas" in linha_lower:
            coletando_pessoa_envolvida = True
        

        # Capturar as pessoas envolvidas (incluindo vítima, autor, etc.)
        if coletando_pessoa_envolvida:
            if "vítima" in linha_lower or "comunicante" in linha_lower:
                partes = linha.split("(")
                if len(partes) > 0:
                    pessoa_envolvida = partes[0].strip()
            
            elif "vestígios" in linha_lower:
                coletando_pessoa_envolvida = False  # Parar de coletar ao encontrar "Vestígios"

        # Captura o relato do PC
        if "relato pc:" in linha_lower:
            coletando_relato = True
            partes = linha.split("Relato PC:")
            if len(partes) > 1:
                relato_pc += partes[1].strip() + " "
        elif coletando_relato:
            if "quesitos" in linha_lower:
                coletando_relato = False  # Não interromper o loop ao encontrar "Quesitos"
            relato_pc += linha.strip() + " "
    
    return rai, odin, requisitante, delegacia_requisitante, unidade_afeta, pessoa_envolvida, relato_pc, endereco







def formatar_nome(nome):
    # Lista de preposições e partículas que não devem ser capitalizadas
    preposicoes = ['da', 'de', 'do', 'das', 'dos', 'e']

    palavras = nome.lower().split()
    nome_formatado = [palavras[0].capitalize()]  # Sempre capitalize a primeira palavra

    for palavra in palavras[1:]:
        if palavra in preposicoes:
            nome_formatado.append(palavra)  # Mantém a preposição em minúsculas
        else:
            nome_formatado.append(palavra.capitalize())  # Capitaliza o restante

    return ' '.join(nome_formatado)

# Restante do código permanece o mesmo


def configurar_paragrafo_com_recuo(paragraph, texto, recuo=Cm(1.25)):
    p = paragraph.add_paragraph(texto)
    p.paragraph_format.first_line_indent = recuo  # Recuo na primeira linha
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alinhar o texto como justificado
    return p

def criar_laudo_com_template(dados_pdf, template_path='templates/template_laudo.docx', output_path='laudo_gerado.docx'):
    if not os.path.exists(template_path):
        messagebox.showerror("Erro", "Template de laudo não encontrado.")
        return
    
    # Extrair dados
    rai, odin, requisitante, delegacia_requisitante, unidade_afeta, pessoa_envolvida, relato_pc, endereco = extrair_dados_necessarios(dados_pdf)

    doc = Document(template_path)  # Carregar o template existente

    # Definir a fonte padrão como Gadugi
    style = doc.styles['Normal']
    style.font.name = 'Gadugi'
    style.font.size = Pt(12)

    # Formatação de Título
    doc.add_heading('LAUDO DE PERÍCIA CRIMINAL', level=1).alignment = 1  # Centralizado

    # Subtítulo com linha em branco depois
    p = doc.add_paragraph('EXAME EM LOCAL DE FURTO', style='Title')
    p.alignment = 1  # Centralizado
    doc.add_paragraph('')  # Linha em branco após subtítulo

    # Detalhes do Procedimento
    tabela = doc.add_table(rows=8, cols=2)
    
    # Remover bordas da tabela
    tbl = tabela._tbl
    for cell in tbl.iter_tcs():
        tc_pr = cell.get_or_add_tcPr()
        tc_borders = parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
        tc_pr.append(tc_borders)

    # Ajustar a largura das colunas e alinhar com o texto
    tabela.alignment = 0  # Alinhar tabela à esquerda
    for row in tabela.rows:
        row.cells[0].width = Pt(200)  # Ajustar o tamanho da primeira coluna
        row.cells[1].width = Pt(300)  # Ajustar o tamanho da segunda coluna

    # Primeira linha: Procedimento
    tabela.cell(0, 0).text = "Procedimento:"
    tabela.cell(0, 1).text = f"RAI: {rai}    ODIN: {odin}"

    # Segunda linha: Requisitante
    tabela.cell(1, 0).text = "Requisitante:"
    tabela.cell(1, 1).text = requisitante

    # Terceira linha: Delegacia Requisitante
    tabela.cell(2, 0).text = "Delegacia Requisitante:"
    tabela.cell(2, 1).text = delegacia_requisitante

    # Quarta linha: Unidade Afeta
    tabela.cell(3, 0).text = "Unidade Afeta:"
    tabela.cell(3, 1).text = unidade_afeta

    # Quinta linha: Pessoa Envolvida
    tabela.cell(4, 0).text = "Pessoa Envolvida:"
    tabela.cell(4, 1).text = formatar_nome(pessoa_envolvida) + " (Vítima)"



    # Sexta linha: Unidade Pericial
    tabela.cell(5, 0).text = "Unidade Pericial:"
    tabela.cell(5, 1).text = "Seção de Perícias Externas da 3ªCRPTC/Formosa."

    # Sétima linha: Perito Criminal
    tabela.cell(6, 0).text = "Perito Criminal:"
    tabela.cell(6, 1).text = "COLOCAR O NOME DO PERITO AQUI"

    # Oitava linha: Início do Exame
    tabela.cell(7, 0).text = "Início do Exame:"
    tabela.cell(7, 1).text = "16 de janeiro de 2023."

    # Ajustando a formatação da tabela
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragrafo in celula.paragraphs:
                paragrafo.style.font.size = Pt(11)
    

    # Adicionar subtítulos com linha em branco antes
    doc.add_paragraph('')  # Linha em branco antes de 1. HISTÓRICO
    doc.add_heading('1. HISTÓRICO', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, relato_pc)
    doc.add_paragraph('')  # Linha em branco antes de 2. OBJETIVO DA PERÍCIA

    doc.add_heading('2. OBJETIVO DA PERÍCIA', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "O objetivo do presente trabalho pericial é estabelecer a materialidade dos fatos declarados sobre a ocorrência, buscando os elementos comprobatórios e os meios e/ou instrumentos utilizados na perpetração do ato, bem como sua autoria.")

     # Seção 3.isolamento
    doc.add_paragraph('')  # Linha em branco 
    doc.add_heading('3. ISOLAMENTO E PRESERVAÇÃO', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "Segundo o Código de Processo Penal – CPP:")
    configurar_paragrafo_com_recuo(doc, "a) Art. 6º - Logo que tiver conhecimento da prática da infração penal, a autoridade policial deverá:")
    configurar_paragrafo_com_recuo(doc, "        I - dirigir-se ao local, providenciando para que não se alterem o estado e conservação das coisas, até a chegada dos peritos criminais;")
    configurar_paragrafo_com_recuo(doc, "        II - apreender os objetos que tiverem relação com o fato, após liberados pelos peritos criminais;")
    configurar_paragrafo_com_recuo(doc, " b) Art. 158-C. § 2º É proibida a entrada em locais isolados bem como a remoção de quaisquer vestígios de locais de crime antes da liberação por parte do perito responsável, sendo tipificada como fraude processual a sua realização.")
    configurar_paragrafo_com_recuo(doc, "c) Art. 169. Parágrafo único. Os peritos registrarão, no laudo, as alterações do estado das coisas e discutirão, no relatório, as consequências dessas alterações na dinâmica dos fatos.")
    configurar_paragrafo_com_recuo(doc, "Neste sentido, o local foi isolado e preservado, conforme as normas legais vigentes, para a realização do exame pericial.")
    configurar_paragrafo_com_recuo(doc, "") # Linha em branco
    # Seção 4 DESCRIÇÃO DO LOCAL
    # Seção 4.1 Endereçament
    doc.add_heading('4. DESCRIÇÃO', level=2).alignment = 0
    doc.add_heading('4.1 Endereçamento', level=2).alignment = 0  # Subtítulo de nível 
    configurar_paragrafo_com_recuo(doc, endereco)  # Adiciona o endereço capturad
    doc.add_heading('4.2 Do local mediato/imediato', level=2).alignment = 0  # Subtítulo de nível 
    configurar_paragrafo_com_recuo(doc, "Tratava-se de um imóvel edificado em alvenaria, ao nível do solo, protegido por muros em todo seu perímetro, sendo que sobre o muro havia a proteção de cerca elétrica. O referido imóvel estava situado em região predominantemente residencial e contava com boa infraestrutura do poder público, como calçamento, asfaltamento, água encanada e luz elétrica, inclusive iluminação pública.")  # Adiciona o endereço capturado
    doc.add_paragraph('')  # Linha em branco antes de 4.2 DO LOCAL MEDIATO/IMEDIATO

  
    doc.add_heading('5. EXAMES', level=2).alignment = 0
    doc.add_heading('5.1 No imóvel', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "Vistoriando-se minunciosamente o imóvel, constatou-se:")
    doc.add_paragraph('')  # Linha em branco antes de 5.1 No imóvel
    doc.add_heading('5.2 Pesquisa por impressões papilares', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "Foram realizadas buscas por impressões papilares nas superfícies que poderiam ter sido tocadas pelo(s) autore(s), contudo não foram constatados fragmentos de impressões papilares.")
    doc.add_heading('5.3 Pesquisa por material biológico', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "Durante processamento do local, não foram observados vestígios que indicassem a presença de material biológico passível de pesquisa por DNA.")

    doc.add_paragraph('')  # Linha em branco antes de 6. Exames Complementares
    doc.add_heading('6. EXAMES COMPLEMENTARES', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "Exames complementares de Papiloscopia não solicitados devido ausência de vestígios.")
    configurar_paragrafo_com_recuo(doc, "Exames complementares de DNA não solicitados devido ausência de vestígios.")

    doc.add_paragraph('')  # Linha em branco antes de 7. DISCUSSÃO TÉCNICA/CONCLUSÃO
    doc.add_heading('7. DISCUSSÃO TÉCNICA/CONCLUSÃO', level=2).alignment = 0
    configurar_paragrafo_com_recuo(doc, "Após a análise minuciosa do local e dos vestígios encontrados, tecemos algumas considerações a respeito do que fora visto e examinado: não foram observados vestígios que indicassem arrombamento, contudo, da forma que a janela do quarto ficava aberta para que o cabo de antena pudesse passar para dentro do imóvel permitiria que a casa fosse invadida sem a necessidade de arrombá-la.")
    configurar_paragrafo_com_recuo(doc, "Por fim, consideramos que a completa elucidação do presente caso fica vinculada à correlação dos elementos materiais extrínsecos inerentes ao fato, acima expostos, com outros meios de prova, porventura existentes.")
    configurar_paragrafo_com_recuo(doc, "Nada mais havendo a considerar, o Perito Criminal encerra o presente laudo pericial ilustrado com 7 fotografia(s), realizada(s) pelo próprio perito relator.")
    configurar_paragrafo_com_recuo(doc, "É o que se tem a relatar.")

    doc.add_paragraph('')
    doc.add_paragraph('')  # Linha em branco antes da data
    p = configurar_paragrafo_com_recuo(doc, "COLOCAR CIDADE AQUI, COLOCAR DATA DE HOJE AQUI.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhar o texto à direita
    doc.add_paragraph('')  # Linha em branco antes da assinatura
    p = configurar_paragrafo_com_recuo(doc, "COLOCAR O NOME DO PERITO AQUI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinhar o texto ao centro
    p = configurar_paragrafo_com_recuo(doc, "Perito Criminal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinhar o texto ao centro


    # Salvar o laudo gerado
    doc.save(output_path)
    messagebox.showinfo("Sucesso", f"Laudo gerado: {output_path}")

def selecionar_pdf():
    pdf_path = filedialog.askopenfilename(
        title="Selecione o arquivo de ocorrência",
        filetypes=(("Arquivos PDF", "*.pdf"),)
    )
    if pdf_path:
        try:
            dados_extraidos = ler_ocorrencia(pdf_path)
            output_path = os.path.join(os.path.dirname(pdf_path), 'laudo_gerado.docx')
            criar_laudo_com_template(dados_extraidos, output_path=output_path)
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro: {str(e)}")

# Interface gráfica com Tkinter
def criar_interface():
    root = tk.Tk()
    root.title("Gerador de Laudos")
    root.geometry("600x300")
    btn_selecionar_pdf = tk.Button(root, text="Selecionar Ocorrência (PDF)", command=selecionar_pdf)
    btn_selecionar_pdf.pack(pady=20)
    root.mainloop()

if __name__ == "__main__":
    criar_interface()